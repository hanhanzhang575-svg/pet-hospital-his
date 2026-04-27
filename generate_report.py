# -*- coding: utf-8 -*-
"""生成宠物医院系统分析报告（Word格式）"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_system_analysis_report():
    """创建系统分析报告"""
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ==================== 第一部分：封面和基本信息 ====================
    title = doc.add_heading('白之助宠物医院信息管理系统', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('课程大作业：系统分析报告', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息表
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    rows_data = [
        ("系统名称", "白之助宠物医院信息管理系统 (Shironosuke Pet Hospital)"),
        ("系统类型", "宠物医疗机构管理信息系统"),
        ("技术框架", "FastAPI + Vue 3 + SQLAlchemy + SQLite"),
        ("开发团队", "信息系统开发组"),
        ("开发时间", "2024年"),
        ("报告日期", datetime.now().strftime("%Y年%m月%d日")),
    ]
    
    for idx, (key, value) in enumerate(rows_data):
        cells = info_table.rows[idx].cells
        cells[0].text = key
        cells[1].text = value
    
    doc.add_paragraph()
    
    # ==================== 第二部分：系统概述 ====================
    doc.add_heading('一、系统概述', level=1)
    
    doc.add_heading('1.1 管理目标', level=2)
    goals = doc.add_paragraph(style='List Bullet')
    goals.add_run('提高宠物医疗服务质量：').bold = True
    goals.add_run('通过规范化的档案管理、预约挂号、病历记录，确保诊疗过程的科学性和可追溯性')
    
    goals = doc.add_paragraph(style='List Bullet')
    goals.add_run('优化资源配置效率：').bold = True
    goals.add_run('实现医生排班、笑舍分配、药品库存的智能管理，降低运营成本')
    
    goals = doc.add_paragraph(style='List Bullet')
    goals.add_run('增强患者体验：').bold = True
    goals.add_run('为主人提供便捷的预约、查询、支付等自助服务')
    
    goals = doc.add_paragraph(style='List Bullet')
    goals.add_run('支持医疗决策：').bold = True
    goals.add_run('通过知识图谱和AI辅助诊断，辅助医生进行科学决策')
    
    goals = doc.add_paragraph(style='List Bullet')
    goals.add_run('建立数据基础：').bold = True
    goals.add_run('收集和管理宠物医疗大数据，为后续分析和研究提供支撑')
    
    # ==================== 第三部分：系统结构 ====================
    doc.add_heading('1.2 系统结构', level=2)
    
    doc.add_paragraph('系统采用典型的三层架构模式：', style='Normal')
    
    structure_items = [
        ('表现层（Presentation Layer）', '基于Vue 3框架的前端界面，为不同角色（医生、护士、主人、管理员等）提供定制化的用户界面'),
        ('业务逻辑层（Business Logic Layer）', '基于FastAPI框架实现的后端服务，包含核心业务逻辑、权限验证、工作流控制等'),
        ('数据访问层（Data Access Layer）', '基于SQLAlchemy ORM框架，访问SQLite数据库，实现数据持久化'),
    ]
    
    for layer, description in structure_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(layer).bold = True
        p.add_run('：' + description)
    
    doc.add_paragraph('分层设计的优势：', style='Normal')
    advantages = [
        '清晰的职责分工，易于开发和维护',
        '接口标准化，支持系统扩展',
        '底层数据变化不影响表现层',
        '便于单元测试和集成测试',
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    # ==================== 第四部分：功能模块 ====================
    doc.add_heading('二、功能、处理、流程', level=1)
    
    doc.add_heading('2.1 核心功能模块', level=2)
    
    modules_data = [
        ('用户与权限管理', '系统用户包括兽医、护士、出纳、管理员等，每个角色有不同权限。支持基于角色的访问控制（RBAC）。'),
        ('主人与宠物档案', '记录宠物主人的基本信息、联系方式，以及宠物的品种、性别、出生日期、过敏史等健康档案。'),
        ('门诊预约挂号', '主人通过线上或线下预约，系统自动分配医生和时间段。支持紧急插号和优先级管理。'),
        ('就诊与病历管理', '记录宠物的主诉、体检结果、诊断结果、治疗方案，为每次就诊生成详细的病历。'),
        ('处方管理', '医生开具处方，系统自动计算所需药品数量，冻结库存，防止超量。处方设置有效期（2小时内必须缴费）。'),
        ('药品库存管理', '维护药品主数据（名称、规格、价格等），按分院跟踪库存。支持安全库存预警。'),
        ('住院与笑舍管理', '分配笑舍单元给住院宠物，记录笑舍状态。支持相邻笑舍的传染病隔离策略。'),
        ('护理日志管理', '护士定时记录住院宠物的体温、心率、护理备注，建立护理历史。'),
        ('收费与发票管理', '记录各类收费项目（诊疗费、药品费、住院费等），生成发票，支持支付结算。'),
        ('AI诊断辅助', '调用医学知识图谱，为医生提供诊断建议。记录医生与AI的决策偏差，实现人机决策审计。'),
    ]
    
    for module, desc in modules_data:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(module).bold = True
        p.add_run('：' + desc)
    
    doc.add_heading('2.2 主要业务流程', level=2)
    
    processes = [
        ('预约挂号流程', [
            '1. 主人通过App/Web发起预约申请',
            '2. 系统检查医生时间表，计算优先级',
            '3. 自动分配医生和时间槽位',
            '4. 主人收到预约确认',
            '5. 到院时，前台确认预约信息',
        ]),
        ('就诊流程', [
            '1. 医生查看预约列表和宠物档案',
            '2. 进行体检，记录主诉和体检结果',
            '3. 调用AI辅助诊断建议',
            '4. 医生确认诊断、制定治疗方案',
            '5. 系统生成病历和处方',
            '6. 主人完成支付',
        ]),
        ('处方与调配流程', [
            '1. 医生开具处方，明确剂量、用法、疗程',
            '2. 系统自动计算所需药品数量',
            '3. 冻结药品库存（锁定库存）',
            '4. 药房工作人员查看待调配处方',
            '5. 进行药品配对、核实、打印标签',
            '6. 主人领取药品（2小时内必须领取）',
        ]),
        ('住院管理流程', [
            '1. 医生下达住院指令',
            '2. 系统分配合适的笑舍单元',
            '3. 记录住院时间、保证金',
            '4. 护士定时记录护理日志（体温、心率等）',
            '5. 医生日常查房、调整治疗方案',
            '6. 出院时结清费用',
        ]),
    ]
    
    for process_name, steps in processes:
        doc.add_heading(process_name, level=3)
        for step in steps:
            doc.add_paragraph(step, style='List Bullet')
    
    # ==================== 第五部分：组成要素 ====================
    doc.add_heading('三、系统组成要素', level=1)
    
    doc.add_heading('3.1 核心数据实体', level=2)
    
    entities = [
        ('用户(User)', '系统操作员，包括医生、护士、出纳、管理员等'),
        ('角色(Role)', '权限集合，定义不同用户的操作权限'),
        ('主人(Owner)', '宠物所有者，记录联系信息和会员等级'),
        ('宠物(Pet)', '患者主体，记录品种、性别、出生日期、过敏史等'),
        ('预约(Appointment)', '挂号记录，绑定医生、宠物、时间'),
        ('病历(MedicalRecord)', '就诊记录，包含主诉、体检、诊断、治疗方案'),
        ('诊断(Diagnosis)', '诊断结果，可能是单个或多个诊断'),
        ('处方(Prescription)', '药物治疗方案，包含多个处方项目'),
        ('药品(Drug)', '药品主数据，包含名称、规格、价格等'),
        ('药品库存(DrugInventory)', '分院级别的药品库存'),
        ('笑舍(CageUnit)', '住院笑舍单元，记录状态和当前住户'),
        ('住院记录(InpatientRecord)', '住院信息，绑定宠物、笑舍、医生'),
        ('护理日志(NursingLog)', '护理记录，包含体温、心率、护理备注'),
        ('发票(Invoice)', '收费单据，聚合各类收费项目'),
    ]
    
    for entity, description in entities:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(entity).bold = True
        p.add_run('：' + description)
    
    doc.add_heading('3.2 接口与传输格式', level=2)
    
    interfaces = [
        ('RESTful API', '所有后端服务通过标准HTTP REST接口暴露，前端通过JSON与后端通信'),
        ('认证机制', '使用JWT（JSON Web Token）进行无状态身份验证'),
        ('异常处理', '统一的错误响应格式，包含状态码、错误消息、错误详情'),
        ('数据验证', '请求数据通过Pydantic schemas进行验证'),
    ]
    
    for interface, desc in interfaces:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(interface).bold = True
        p.add_run('：' + desc)
    
    # ==================== 第六部分：要素间关系 ====================
    doc.add_heading('四、接口（要素间关系）', level=1)
    
    doc.add_heading('4.1 核心关系图', level=2)
    
    relationships = [
        ('用户 --> 宠物', '医生通过预约与宠物关联，进行诊疗'),
        ('主人 --> 宠物', '一对多关系，一个主人可拥有多个宠物'),
        ('宠物 --> 预约', '一对多关系，一个宠物可有多个预约'),
        ('预约 --> 医生(User)', '多对一关系，每个预约分配给一个医生'),
        ('预约 --> 病历', '一对一关系，一个预约对应一份病历'),
        ('病历 --> 诊断', '一对多关系，一份病历可有多个诊断'),
        ('病历 --> 处方', '一对多关系，一份病历可开多个处方'),
        ('处方 --> 处方项目', '一对多关系，一个处方包含多个药品'),
        ('处方项目 --> 药品', '多对一关系，多个处方项目可引用同一药品'),
        ('药品 --> 库存', '一对多关系，一个药品在多个分院有库存'),
        ('宠物 --> 住院记录', '一对多关系，一个宠物可多次住院'),
        ('住院记录 --> 笑舍', '多对一关系，多个住院记录可使用同一笑舍（不同时间）'),
        ('住院记录 --> 护理日志', '一对多关系，一条住院记录有多条护理日志'),
        ('主人 --> 发票', '一对多关系，一个主人可有多个发票'),
        ('发票 --> 发票项目', '一对多关系，一个发票包含多个收费项目'),
    ]
    
    for rel, description in relationships:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(rel).bold = True
        p.add_run('：' + description)
    
    doc.add_heading('4.2 数据流关键路径', level=2)
    
    data_flows = [
        ('预约 --> 病历 --> 诊断 --> 处方', '患者从预约到就诊到开药的完整数据流'),
        ('处方 --> 药品库存 --> 发票', '药品供应链和费用结算的数据流'),
        ('宠物档案 --> 就诊历史 --> 统计分析', '患者数据的累积和分析流程'),
        ('住院指令 --> 笑舍分配 --> 护理日志 --> 出院结算', '住院全生命周期的数据流'),
    ]
    
    for flow, description in data_flows:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(flow).bold = True
        p.add_run('：' + description)
    
    # ==================== 第七部分：系统边界 ====================
    doc.add_heading('五、系统边界', level=1)
    
    doc.add_heading('5.1 系统范围内的业务', level=2)
    
    inside = [
        '宠物主人和宠物档案管理',
        '预约挂号和时间管理',
        '就诊、诊疗、病历记录',
        '处方开具和药品调配',
        '住院和笑舍管理',
        '护理日志记录',
        '收费和发票管理',
        'AI诊断辅助和知识图谱查询',
        '用户权限和访问控制',
        '系统性能监控和健康检查',
    ]
    
    for item in inside:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('5.2 系统范围外的业务', level=2)
    
    outside = [
        '财务核算和报表（仅记录费用，不做会计处理）',
        '医疗设备管理（如超声、X光等设备）',
        '物资采购和供应商管理',
        '人力资源和薪酬管理',
        '医学科研和论文管理',
        '医疗保险对接（与第三方系统对接）',
        '监管部门上报（数据导出接口预留）',
    ]
    
    for item in outside:
        doc.add_paragraph(item, style='List Bullet')
    
    # ==================== 第八部分：信息需求与输出 ====================
    doc.add_heading('六、输入与输出', level=1)
    
    doc.add_heading('6.1 主要输入信息需求', level=2)
    
    p = doc.add_heading('宠物主人端输入', level=3)
    inputs_owner = [
        '宠物信息（品种、性别、体重、过敏史等）',
        '预约申请（首选医生、就诊时间、主诉症状）',
        '就诊反馈（用药效果、用药后反应）',
        '支付信息（支付方式、发票寄送地址）',
    ]
    for inp in inputs_owner:
        doc.add_paragraph(inp, style='List Bullet')
    
    p = doc.add_heading('医生/护士端输入', level=3)
    inputs_medical = [
        '就诊记录（主诉、体检结果、初步诊断）',
        '诊疗决策（最终诊断、诊断代码、严重程度）',
        '处方信息（药品、剂量、用法、疗程）',
        '住院指令（住院原因、预期疗程、特殊说明）',
        '护理日志（体温、心率、排尿排便、特殊护理等）',
        '出院总结（治疗效果、出院建议、复诊时间）',
    ]
    for inp in inputs_medical:
        doc.add_paragraph(inp, style='List Bullet')
    
    p = doc.add_heading('管理员端输入', level=3)
    inputs_admin = [
        '药品信息（新增药品、更新价格）',
        '人员信息（员工档案、权限分配）',
        '系统配置（预约时间、笑舍分布、分院信息）',
        '黑名单管理（失信用户、异常行为记录）',
    ]
    for inp in inputs_admin:
        doc.add_paragraph(inp, style='List Bullet')
    
    doc.add_heading('6.2 主要输出与管理决策信息', level=2)
    
    p = doc.add_heading('面向宠物主人的输出', level=3)
    outputs_owner = [
        '预约确认（时间、医生、排队位置）',
        '就诊报告（诊断结果、治疗方案、用药指导）',
        '药品配备单（药品清单、用法用量、服用周期）',
        '收费发票（明细、支付状态、电子证明）',
        '健康档案（过去就诊记录、用药史、预防接种记录）',
    ]
    for out in outputs_owner:
        doc.add_paragraph(out, style='List Bullet')
    
    p = doc.add_heading('面向医疗团队的输出', level=3)
    outputs_medical = [
        '医生工作列表（待诊患者、预约时间表、待办任务）',
        '宠物档案与历史（完整病历、过往诊疗、用药记录）',
        'AI诊断建议（知识图谱推理、诊断置信度、参考文献）',
        '库存和处方状态（药品库存预警、待调配处方、逾期处方）',
        '住院患者列表（在院宠物、笑舍使用、护理需求）',
        '个人统计（日均就诊数、开单总额、诊疗准确率）',
    ]
    for out in outputs_medical:
        doc.add_paragraph(out, style='List Bullet')
    
    p = doc.add_heading('面向管理层的输出', level=3)
    outputs_mgmt = [
        '运营仪表板（日/周/月统计、科室对比、收入趋势）',
        '医疗质量指标（诊疗准确率、患者满意度、医患纠纷比率）',
        '资源利用率（医生排班满度、笑舍床位利用率、药品周转率）',
        '财务分析报告（收入、成本、利润率、应收账款）',
        '预警告警（库存预警、逾期处方、异常费用、系统故障）',
    ]
    for out in outputs_mgmt:
        doc.add_paragraph(out, style='List Bullet')
    
    # ==================== 第九部分：系统控制 ====================
    doc.add_heading('七、系统控制', level=1)
    
    doc.add_heading('7.1 访问控制', level=2)
    
    controls_access = [
        ('基于角色的访问控制(RBAC)', '每个用户属于一个角色（医生、护士、出纳、管理员），不同角色可访问的功能和数据不同'),
        ('医生', '可查看自己名下的预约、患者档案、处方；不能查看其他医生的记录'),
        ('护士', '可记录护理日志、查看住院患者；不能开具诊疗决策'),
        ('出纳', '可查看和处理收费发票；不能修改医疗决策'),
        ('管理员', '可管理用户、配置系统、查看所有数据'),
        ('宠物主人', '仅可查看自己宠物的预约、病历摘要、账单'),
    ]
    
    for control_name, description in controls_access:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(control_name).bold = True
        p.add_run('：' + description)
    
    doc.add_heading('7.2 数据完整性控制', level=2)
    
    controls_integrity = [
        ('病历不可追溯修改', '病历一旦创建不可删除，修改需要记录原值和修改原因（审计日志）'),
        ('处方有效期控制', '处方必须在指定时间内（2小时）支付缴费，过期自动作废'),
        ('库存冻结机制', '处方一旦生成，立即冻结所需药品库存，防止超量或重复调配'),
        ('笑舍隔离策略', '传染病宠物不能住在相邻笑舍，系统自动检查'),
        ('费用核对', '发票项目和金额必须与实际医疗记录对应，防止漏收或多收'),
    ]
    
    for control_name, description in controls_integrity:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(control_name).bold = True
        p.add_run('：' + description)
    
    doc.add_heading('7.3 业务规则控制', level=2)
    
    controls_business = [
        ('预约分配规则', '考虑医生当日排班、可用时段、患者等级、紧急程度等因素自动分配'),
        ('优先级评分', '根据紧急程度、等候时间、患者等级等因素计算动态优先级'),
        ('药品库存预警', '库存低于安全库存时生成预警，提醒采购'),
        ('超期处理', '逾期未缴费的处方和预约自动作废，同时记录异常'),
        ('AI决策审计', '对医生采纳或拒绝AI建议的情况进行追踪，异常偏差触发告警'),
    ]
    
    for control_name, description in controls_business:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(control_name).bold = True
        p.add_run('：' + description)
    
    doc.add_heading('7.4 系统安全控制', level=2)
    
    controls_security = [
        ('密码策略', '密码必须包含大小写字母、数字、特殊字符，最小长度8位'),
        ('会话管理', '使用JWT令牌，自动过期机制，支持强制登出'),
        ('操作日志', '所有用户操作记录详细日志，包括操作人、操作时间、操作内容'),
        ('数据备份', '定期备份数据库，确保数据安全和可恢复性'),
        ('加密传输', 'API通信支持HTTPS/TLS加密'),
    ]
    
    for control_name, description in controls_security:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(control_name).bold = True
        p.add_run('：' + description)
    
    # ==================== 第十部分：数据库表结构 ====================
    doc.add_heading('八、数据库表结构与SQL关联', level=1)
    
    doc.add_heading('8.1 核心表设计', level=2)
    
    db_tables = [
        ('roles', [
            ('id', 'INTEGER PRIMARY KEY', '角色ID'),
            ('name', 'VARCHAR(50) UNIQUE', '角色名称'),
            ('description', 'VARCHAR(255)', '角色描述'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('users', [
            ('id', 'INTEGER PRIMARY KEY', '用户ID'),
            ('employee_code', 'VARCHAR(20) UNIQUE', '员工编号'),
            ('username', 'VARCHAR(50) UNIQUE', '用户名'),
            ('password_hash', 'VARCHAR(255)', '密码哈希'),
            ('full_name', 'VARCHAR(100)', '真实姓名'),
            ('role_id', 'INTEGER FK', '角色ID'),
            ('branch_code', 'VARCHAR(10)', '分院代码'),
            ('is_active', 'BOOLEAN', '是否激活'),
            ('is_licensed_vet', 'BOOLEAN', '是否持证兽医'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('owners', [
            ('id', 'INTEGER PRIMARY KEY', '主人ID'),
            ('owner_code', 'VARCHAR(20) UNIQUE', '客户编号'),
            ('name', 'VARCHAR(100)', '主人名字'),
            ('phone', 'VARCHAR(20)', '手机号'),
            ('telephone', 'VARCHAR(20)', '座机号'),
            ('id_card', 'VARCHAR(30)', '身份证号'),
            ('address', 'VARCHAR(255)', '地址'),
            ('member_level', 'VARCHAR(20)', '会员等级'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('pets', [
            ('id', 'INTEGER PRIMARY KEY', '宠物ID'),
            ('pet_code', 'VARCHAR(20) UNIQUE', '宠物编号'),
            ('name', 'VARCHAR(100)', '宠物名称'),
            ('species', 'VARCHAR(20)', '物种'),
            ('breed', 'VARCHAR(50)', '品种'),
            ('gender', 'VARCHAR(10)', '性别'),
            ('birth_date', 'DATE', '出生日期'),
            ('color', 'VARCHAR(100)', '颜色/花纹'),
            ('weight', 'FLOAT', '体重(kg)'),
            ('allergy_history', 'JSON', '过敏史'),
            ('owner_id', 'INTEGER FK', '主人ID'),
            ('clinic_id', 'VARCHAR(10)', '就诊分院'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('appointments', [
            ('id', 'INTEGER PRIMARY KEY', '预约ID'),
            ('record_code', 'VARCHAR(30) UNIQUE', '预约号'),
            ('pet_id', 'INTEGER FK', '宠物ID'),
            ('doctor_id', 'INTEGER FK', '医生ID'),
            ('clinic_id', 'VARCHAR(10)', '分院代码'),
            ('scheduled_time', 'DATETIME', '预约时间'),
            ('urgency_level', 'VARCHAR(20)', '紧急程度'),
            ('status', 'VARCHAR(20)', '预约状态'),
            ('priority_score', 'FLOAT', '优先级评分'),
            ('max_capacity', 'INTEGER', '最大容纳数'),
            ('is_leave', 'BOOLEAN', '是否请假'),
            ('schedule_note', 'VARCHAR(255)', '备注'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('medical_records', [
            ('id', 'INTEGER PRIMARY KEY', '病历ID'),
            ('record_no', 'VARCHAR(30) UNIQUE', '病历号'),
            ('appointment_id', 'INTEGER FK', '预约ID'),
            ('pet_id', 'INTEGER FK', '宠物ID'),
            ('vet_id', 'INTEGER FK', '兽医ID'),
            ('chief_complaint', 'TEXT', '主诉'),
            ('exam_notes', 'TEXT', '体检记录'),
            ('diagnosis', 'VARCHAR(100)', '诊断'),
            ('treatment_plan', 'TEXT', '治疗方案'),
            ('is_voided', 'BOOLEAN', '是否作废'),
            ('kg_evidence_id', 'VARCHAR(100)', '知识图谱关联'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('prescriptions', [
            ('id', 'INTEGER PRIMARY KEY', '处方ID'),
            ('prescription_code', 'VARCHAR(30) UNIQUE', '处方号'),
            ('medical_record_id', 'INTEGER FK', '病历ID'),
            ('doctor_id', 'INTEGER FK', '开方医生'),
            ('status', 'VARCHAR(20)', '处方状态'),
            ('expire_at', 'DATETIME', '处方过期时间'),
            ('locked_inventory', 'JSON', '冻结库存'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('drugs', [
            ('id', 'INTEGER PRIMARY KEY', '药品ID'),
            ('drug_code', 'VARCHAR(30) UNIQUE', '药品编号'),
            ('name', 'VARCHAR(100)', '药品名称'),
            ('dosage_form', 'VARCHAR(20)', '剂型'),
            ('category', 'VARCHAR(50)', '分类'),
            ('unit', 'VARCHAR(20)', '单位'),
            ('unit_price', 'FLOAT', '单价'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('drug_inventory', [
            ('id', 'INTEGER PRIMARY KEY', '库存ID'),
            ('drug_id', 'INTEGER FK', '药品ID'),
            ('branch_code', 'VARCHAR(10)', '分院代码'),
            ('stock_qty', 'FLOAT', '库存数量'),
            ('safety_stock', 'FLOAT', '安全库存'),
            ('frozen_for_prescription', 'BOOLEAN', '是否冻结'),
            ('expiry_date', 'DATETIME', '过期日期'),
            ('updated_at', 'DATETIME', '更新时间'),
        ]),
        ('cage_units', [
            ('id', 'INTEGER PRIMARY KEY', '笑舍ID'),
            ('cage_code', 'VARCHAR(20) UNIQUE', '笑舍编号'),
            ('clinic_id', 'VARCHAR(10)', '分院代码'),
            ('zone_type', 'VARCHAR(20)', '区域类型'),
            ('status', 'VARCHAR(20)', '笑舍状态'),
            ('current_pet_id', 'INTEGER FK', '当前宠物'),
            ('adjacent_cage_ids', 'JSON', '相邻笑舍'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
        ('inpatient_records', [
            ('id', 'INTEGER PRIMARY KEY', '住院ID'),
            ('pet_id', 'INTEGER FK', '宠物ID'),
            ('cage_id', 'INTEGER FK', '笑舍ID'),
            ('doctor_id', 'INTEGER FK', '主治医生'),
            ('clinic_id', 'VARCHAR(10)', '分院代码'),
            ('admission_time', 'DATETIME', '入院时间'),
            ('discharge_time', 'DATETIME', '出院时间'),
            ('deposit_amount', 'FLOAT', '预存金额'),
            ('consumed_amount', 'FLOAT', '消费金额'),
            ('status', 'VARCHAR(20)', '住院状态'),
            ('created_at', 'DATETIME', '创建时间'),
        ]),
    ]
    
    for table_name, columns in db_tables:
        p = doc.add_heading(f'表：{table_name}', level=3)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        header_cells = table.rows[0].cells
        header_cells[0].text = '字段名'
        header_cells[1].text = '数据类型'
        header_cells[2].text = '说明'
        
        for col_name, col_type, col_desc in columns:
            row_cells = table.add_row().cells
            row_cells[0].text = col_name
            row_cells[1].text = col_type
            row_cells[2].text = col_desc
    
    doc.add_heading('8.2 表间主外键关系', level=2)
    
    fk_relations = [
        ('users.role_id --> roles.id', '用户属于某个角色'),
        ('pets.owner_id --> owners.id', '宠物属于某个主人'),
        ('appointments.pet_id --> pets.id', '预约关联某个宠物'),
        ('appointments.doctor_id --> users.id', '预约分配给某个医生'),
        ('medical_records.appointment_id --> appointments.id', '病历对应某个预约'),
        ('medical_records.pet_id --> pets.id', '病历记录某个宠物'),
        ('medical_records.vet_id --> users.id', '病历由某个兽医创建'),
        ('prescriptions.medical_record_id --> medical_records.id', '处方对应某份病历'),
        ('prescriptions.doctor_id --> users.id', '处方由某个医生开具'),
        ('prescription_items.prescription_id --> prescriptions.id', '处方项目属于某个处方'),
        ('prescription_items.drug_id --> drugs.id', '处方项目指定某种药品'),
        ('drug_inventory.drug_id --> drugs.id', '库存记录某种药品'),
        ('cage_units.current_pet_id --> pets.id', '笑舍当前住宿的宠物'),
        ('inpatient_records.pet_id --> pets.id', '住院记录某个宠物'),
        ('inpatient_records.cage_id --> cage_units.id', '住院记录分配的笑舍'),
        ('inpatient_records.doctor_id --> users.id', '住院由某个医生主治'),
        ('invoices.owner_id --> owners.id', '发票对应某个主人'),
        ('invoices.appointment_id --> appointments.id', '发票关联某个预约'),
    ]
    
    for fk, description in fk_relations:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(fk).bold = True
        p.add_run('：' + description)
    
    # ==================== 总结 ====================
    doc.add_heading('九、系统总结', level=1)
    
    summary = doc.add_paragraph(
        '本系统是一个功能完整的宠物医院信息管理系统，涵盖患者档案、预约挂号、诊疗记录、药品管理、'
        '住院照护、收费结算等核心业务。系统采用现代化的三层架构，引入AI辅助诊断，实现了数据驱动的'
        '医疗管理。通过严格的访问控制、数据完整性约束、业务规则控制，确保了诊疗质量和患者安全。'
    )
    
    doc.add_paragraph()
    
    closing = doc.add_paragraph(
        '系统具有良好的可扩展性和集成性，为今后集成医疗保险、远程诊疗、科研分析等功能奠定了坚实基础。'
    )
    
    # 保存文档
    output_path = r'C:\Users\zhangmohan\Desktop\PetHospital_Project\信息系统\白之助宠物医院系统分析报告.docx'
    doc.save(output_path)
    print(f"Report generated: {output_path}")

if __name__ == '__main__':
    create_system_analysis_report()
